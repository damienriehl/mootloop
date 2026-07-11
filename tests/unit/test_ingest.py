"""Unit tests for corpus models and folder ingestion."""

from __future__ import annotations

from pathlib import Path

import pytest

from mootloop.errors import IngestError
from mootloop.ingest import ingest_folder, set_doc_tag
from mootloop.models.corpus import DocRole, Manifest
from mootloop.vault import create_vault
from tests.conftest import make_matter

NOW = "2026-07-11T12:00:00+00:00"


def _vault(tmp_path: Path) -> Path:
    root = tmp_path / "vault"
    registry = tmp_path / "canaries.json"
    create_vault(root, make_matter(), registry_path=registry)
    return root


def _write(source: Path, name: str, text: str) -> None:
    source.mkdir(parents=True, exist_ok=True)
    (source / name).write_text(text, encoding="utf-8")


def test_ingest_text_passthrough_and_manifest(tmp_path: Path) -> None:
    vault = _vault(tmp_path)
    src = tmp_path / "src"
    _write(src, "notes.txt", "hello world")
    _write(src, "memo.md", "# Memo\n\nbody")

    report = ingest_folder(vault, src, now=NOW)
    assert report.status_counts()["ok"] == 2

    manifest = Manifest.load(vault)
    assert len(manifest.docs) == 2
    for doc in manifest.docs:
        assert doc.doc_id.startswith("doc-")
        assert doc.ingest_status == "ok"
        assert doc.normalized_path is not None
        assert (vault / doc.normalized_path).is_file()
        # original preserved with its extension
        originals = list((vault / "corpus" / "originals").glob(f"{doc.doc_id}.*"))
        assert originals


def test_doc_id_is_content_addressed_and_idempotent(tmp_path: Path) -> None:
    vault = _vault(tmp_path)
    src = tmp_path / "src"
    _write(src, "a.txt", "identical bytes")

    first = ingest_folder(vault, src, now=NOW).entries[0].doc.doc_id
    # Re-ingest the same content under a different name → same content id, no dupes.
    _write(src, "b.txt", "identical bytes")
    ingest_folder(vault, src, now=NOW)
    manifest = Manifest.load(vault)
    assert len(manifest.docs) == 1
    assert manifest.docs[0].doc_id == first


def test_latin1_fallback(tmp_path: Path) -> None:
    vault = _vault(tmp_path)
    src = tmp_path / "src"
    src.mkdir()
    (src / "legacy.txt").write_bytes(b"caf\xe9 na\xefve")  # invalid utf-8

    report = ingest_folder(vault, src, now=NOW)
    doc = report.entries[0].doc
    assert doc.ingest_status == "ok"
    assert doc.normalized_path is not None
    assert "café" in (vault / doc.normalized_path).read_text(encoding="utf-8")


def test_eml_normalizes_headers_as_front_matter(tmp_path: Path) -> None:
    vault = _vault(tmp_path)
    src = tmp_path / "src"
    eml = (
        "From: sender@example.com\n"
        "To: receiver@example.com\n"
        "Subject: Widget shipment\n"
        "Date: Mon, 01 Jun 2026 09:00:00 -0500\n"
        "Content-Type: text/plain\n"
        "\n"
        "The widgets did not arrive on time.\n"
    )
    _write(src, "msg.eml", eml)

    report = ingest_folder(vault, src, now=NOW)
    doc = report.entries[0].doc
    assert doc.ingest_status == "ok"
    assert doc.media_type == "message/rfc822"
    text = (vault / doc.normalized_path).read_text(encoding="utf-8")  # type: ignore[arg-type]
    assert text.startswith("---")
    assert "subject: Widget shipment" in text
    assert "The widgets did not arrive on time." in text


def test_docx_paragraph_extraction(tmp_path: Path) -> None:
    from docx import Document

    vault = _vault(tmp_path)
    src = tmp_path / "src"
    src.mkdir()
    document = Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    document.save(str(src / "brief.docx"))

    report = ingest_folder(vault, src, now=NOW)
    doc = report.entries[0].doc
    assert doc.ingest_status == "ok"
    text = (vault / doc.normalized_path).read_text(encoding="utf-8")  # type: ignore[arg-type]
    assert "First paragraph." in text
    assert "Second paragraph." in text


def test_pdf_needs_conversion(tmp_path: Path) -> None:
    vault = _vault(tmp_path)
    src = tmp_path / "src"
    src.mkdir()
    (src / "scan.pdf").write_bytes(b"%PDF-1.4 fake")

    report = ingest_folder(vault, src, now=NOW)
    doc = report.entries[0].doc
    assert doc.ingest_status == "needs_conversion"
    assert doc.normalized_path is None
    # original still copied so it is not lost
    assert list((vault / "corpus" / "originals").glob(f"{doc.doc_id}.pdf"))


def test_too_large(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    import mootloop.ingest as ingest_mod

    monkeypatch.setattr(ingest_mod, "MAX_BYTES", 8)
    vault = _vault(tmp_path)
    src = tmp_path / "src"
    _write(src, "big.txt", "way more than eight bytes")

    report = ingest_folder(vault, src, now=NOW)
    doc = report.entries[0].doc
    assert doc.ingest_status == "too_large"
    assert doc.normalized_path is None


def test_symlink_fails_closed(tmp_path: Path) -> None:
    vault = _vault(tmp_path)
    src = tmp_path / "src"
    src.mkdir()
    target = tmp_path / "outside.txt"
    target.write_text("secret", encoding="utf-8")
    (src / "link.txt").symlink_to(target)

    report = ingest_folder(vault, src, now=NOW)
    doc = report.entries[0].doc
    assert doc.ingest_status == "unreadable"
    assert "fail closed" in (report.entries[0].reason or "")


def test_hidden_files_skipped(tmp_path: Path) -> None:
    vault = _vault(tmp_path)
    src = tmp_path / "src"
    _write(src, ".hidden.txt", "nope")
    _write(src, "visible.txt", "yes")

    report = ingest_folder(vault, src, now=NOW)
    assert len(report.entries) == 1
    assert report.entries[0].doc.original_name == "visible.txt"


def test_tags_file_applies_role_and_privilege(tmp_path: Path) -> None:
    vault = _vault(tmp_path)
    src = tmp_path / "src"
    _write(src, "complaint.md", "complaint text")
    _write(src, "email1.eml", "From: a@b.co\n\nhi\n")
    _write(src, "email2.eml", "From: c@d.co\n\nprivileged\n")
    tags = tmp_path / "tags.yaml"
    tags.write_text(
        'complaint.md: {role: complaint, privileged: false}\n'
        '"*.eml": {role: correspondence, privileged: false}\n'
        'email2.eml: {role: correspondence, privileged: true}\n',
        encoding="utf-8",
    )

    ingest_folder(vault, src, now=NOW, tags_file=tags)
    manifest = Manifest.load(vault)
    by_name = {d.original_name: d for d in manifest.docs}
    assert by_name["complaint.md"].role == DocRole.COMPLAINT
    assert by_name["email1.eml"].role == DocRole.CORRESPONDENCE
    assert by_name["email1.eml"].privileged is False
    # last matching rule wins → email2 privileged
    assert by_name["email2.eml"].privileged is True


def test_reingest_preserves_prior_manual_tag(tmp_path: Path) -> None:
    vault = _vault(tmp_path)
    src = tmp_path / "src"
    _write(src, "doc.txt", "stable content")
    doc_id = ingest_folder(vault, src, now=NOW).entries[0].doc.doc_id
    set_doc_tag(vault, doc_id, role=DocRole.CLIENT_DOC, privileged=True)

    # Re-ingest with no tags: prior role/privilege must survive.
    ingest_folder(vault, src, now=NOW)
    doc = Manifest.load(vault).get(doc_id)
    assert doc is not None
    assert doc.role == DocRole.CLIENT_DOC
    assert doc.privileged is True


def test_set_doc_tag_unknown_id_raises(tmp_path: Path) -> None:
    vault = _vault(tmp_path)
    with pytest.raises(IngestError):
        set_doc_tag(vault, "doc-deadbeefdeadbeef", role=DocRole.OTHER)


def test_missing_source_dir_raises(tmp_path: Path) -> None:
    vault = _vault(tmp_path)
    with pytest.raises(IngestError):
        ingest_folder(vault, tmp_path / "nope", now=NOW)
