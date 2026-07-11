#!/usr/bin/env python3
"""Generate the two pandoc reference-doc templates (plan D8).

pandoc's ``--reference-doc`` owns the court chrome (page size, margins, base font, a
footer page-number placeholder) that python-docx cannot otherwise inject into rendered
output. We generate the templates once, here, and COMMIT them — they carry NO matter
data, only styling. Two variants:

  - ``generic-mn-district.docx``        — the clean, attested template.
  - ``generic-mn-district-draft.docx``  — adds a "DRAFT — NOT FOR SERVICE" header.

Run: ``uv run python tools/make_reference_docx.py``. Idempotent (overwrites).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

REPO_ROOT = Path(__file__).resolve().parent.parent
COURTS_DIR = REPO_ROOT / "config" / "courts"

# US Letter, 1-inch margins (plan D8).
_LETTER_WIDTH = Inches(8.5)
_LETTER_HEIGHT = Inches(11)
_MARGIN = Inches(1)
_FONT = "Times New Roman"
_FONT_SIZE = Pt(12)


def _build(draft: bool) -> Document:
    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = _FONT_SIZE

    for section in document.sections:
        section.page_width = _LETTER_WIDTH
        section.page_height = _LETTER_HEIGHT
        section.left_margin = section.right_margin = _MARGIN
        section.top_margin = section.bottom_margin = _MARGIN

        # Footer page-number placeholder (static text; pandoc owns the live field).
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = "Page"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if draft:
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = "DRAFT — NOT FOR SERVICE"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header_para.runs[0] if header_para.runs else header_para.add_run()
            run.bold = True

    return document


def main() -> int:
    COURTS_DIR.mkdir(parents=True, exist_ok=True)
    _build(draft=False).save(str(COURTS_DIR / "generic-mn-district.docx"))
    _build(draft=True).save(str(COURTS_DIR / "generic-mn-district-draft.docx"))
    print(f"Wrote reference docs to {COURTS_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
